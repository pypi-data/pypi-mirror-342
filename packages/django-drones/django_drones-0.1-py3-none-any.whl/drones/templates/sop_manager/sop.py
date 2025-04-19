from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Create new SOP document
doc = Document()

# Title Page with branding
doc.add_picture("/mnt/data/AI Logo.png", width=Inches(2))
title = doc.add_heading("Standard Operating Procedures (SOP)", 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
contact_info = doc.add_paragraph("Airborne Images\nTom Stout\nwww.airborne-images.net\n317-987-7387\ntom@airborne-images.net")
contact_info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_page_break()

# SOP Sections
sop_sections = {
    "1. Purpose": "This document outlines the Standard Operating Procedures (SOPs) for Airborne Images, a professional drone service provider specializing in aerial photography, videography, and inspection services.",
    "2. Scope": "These procedures apply to all drone operations conducted by Airborne Images using DJI Mavic 3 Pro aircraft, under FAA Part 107 regulations.",
    "3. Responsibilities": (
        "The Remote Pilot in Command (RPIC) holds ultimate authority and responsibility for the safety and legality of each flight. "
        "All crew members, including Visual Observers (VOs), must follow the direction of the RPIC and be briefed before operations."
    ),
    "4. Equipment": (
        "Primary UAS: DJI Mavic 3 Pro\n"
        "- Ensure all hardware is up-to-date and functioning.\n"
        "- Batteries must be charged, rotated, and logged.\n"
        "- Conduct firmware updates and system calibrations as required."
    ),
    "5. Flight Preparation": (
        "- Check current and forecasted weather.\n"
        "- Review NOTAMs and obtain LAANC authorization if in controlled airspace.\n"
        "- Conduct site survey and risk assessment.\n"
        "- Complete pre-flight checklist and safety briefing with crew."
    ),
    "6. Flight Operations": (
        "- Maintain Visual Line of Sight (VLOS) at all times.\n"
        "- Operate below 400 feet AGL unless otherwise authorized.\n"
        "- Avoid flight over non-participating people and vehicles.\n"
        "- Follow airspace restrictions and comply with ATC coordination.\n"
        "- Announce takeoff and landing verbally to crew."
    ),
    "7. Emergency Procedures": (
        "- Initiate Return-to-Home (RTH) if GPS signal is lost.\n"
        "- Conduct emergency landing in case of loss of control or visual.\n"
        "- Report any incidents or injuries as required by FAA Part 107 regulations."
    ),
    "8. Maintenance and Logging": (
        "- Log all flights, maintenance, and battery cycles in digital or physical logs.\n"
        "- Perform post-flight inspections and maintenance checks.\n"
        "- Replace worn or damaged components before next flight."
    ),
    "9. Training Requirements": (
        "- RPIC must hold a valid FAA Part 107 certification and remain current.\n"
        "- Visual Observers must be trained in basic airspace awareness and visual tracking.\n"
        "- Recurrent training and safety drills must be conducted annually."
    ),
    "10. Document Review": (
        "This SOP is reviewed annually or after significant regulatory updates.\n"
        "All revisions will be documented and distributed to operational personnel."
    )
}

# Add SOP sections to the document
for heading, content in sop_sections.items():
    doc.add_heading(heading, level=1)
    for para in content.split("\n"):
        doc.add_paragraph(para)

# Save the Word document
sop_doc_path = "/mnt/data/Airborne_Images_SOP.docx"
doc.save(sop_doc_path)

sop_doc_path

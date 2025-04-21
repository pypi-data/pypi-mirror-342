import docx
import pytest
import numpy as np
from unittest.mock import patch, MagicMock
from vectoriz.files import FileArgument
from vectoriz.files import FilesFeature


class TestFileArgument:
    def test_add_data_appends_to_lists(self):
        file_arg = FileArgument()
        filename = "test.txt"
        text = "Test content"

        with patch.object(
            FileArgument, "_create_embedding", return_value=np.array([0.1, 0.2, 0.3])
        ):
            file_arg.add_data(filename, text)

            assert file_arg.chunk_names == [filename]
            assert file_arg.text_list == [text]
            assert len(file_arg.embeddings) == 1
            np.testing.assert_array_equal(
                file_arg.embeddings[0], np.array([0.1, 0.2, 0.3])
            )

    def test_add_data_multiple_entries(self):
        file_arg = FileArgument(
            ["existing.txt"], ["existing content"], [np.array([0.5, 0.5, 0.5])]
        )
        filename = "new.txt"
        text = "New content"

        with patch.object(
            FileArgument, "_create_embedding", return_value=np.array([0.7, 0.8, 0.9])
        ):
            file_arg.add_data(filename, text)
            assert file_arg.chunk_names == ["existing.txt", "new.txt"]
            assert file_arg.text_list == ["existing content", "New content"]
            assert len(file_arg.embeddings) == 2
            np.testing.assert_array_equal(
                file_arg.embeddings[1], np.array([0.7, 0.8, 0.9])
            )

    def test_add_data_calls_create_embedding(self):
        file_arg = FileArgument()
        filename = "test.txt"
        text = "Test content"

        with patch.object(FileArgument, "_create_embedding") as mock_create_embedding:
            mock_create_embedding.return_value = np.array([0.1, 0.2, 0.3])
            file_arg.add_data(filename, text)
            mock_create_embedding.assert_called_once_with(text)

    def test_create_embedding_returns_numpy_array(self):
        file_arg = FileArgument()
        text = "Test content"

        with patch("vectoriz.files.TokenTransformer") as mock_transformer:
            mock_instance = mock_transformer.return_value
            mock_instance.text_to_embeddings.return_value = [np.array([0.1, 0.2, 0.3])]

            result = file_arg._create_embedding(text)

            assert isinstance(result, np.ndarray)
            np.testing.assert_array_equal(result, np.array([0.1, 0.2, 0.3]))
            mock_instance.text_to_embeddings.assert_called_once_with([text])

    def test_create_embedding_handles_empty_text(self):
        file_arg = FileArgument()
        text = ""

        with patch("vectoriz.files.TokenTransformer") as mock_transformer:
            mock_instance = mock_transformer.return_value
            mock_instance.text_to_embeddings.return_value = [np.array([0.0, 0.0, 0.0])]

            result = file_arg._create_embedding(text)

            assert isinstance(result, np.ndarray)
            mock_instance.text_to_embeddings.assert_called_once_with([""])

    def test_create_embedding_instantiates_token_transformer(self):
        file_arg = FileArgument()
        text = "Test content"

        with patch("vectoriz.files.TokenTransformer") as mock_transformer:
            mock_instance = mock_transformer.return_value
            mock_instance.text_to_embeddings.return_value = [np.array([0.1, 0.2, 0.3])]

            file_arg._create_embedding(text)

            mock_transformer.assert_called_once()


class TestFilesFeature:
    def test_extract_txt_content_reads_file_correctly(self, tmp_path):
        test_content = "This is test content"
        test_file = tmp_path / "test.txt"
        test_file.write_text(test_content)
        files_feature = FilesFeature()
        result = files_feature._extract_txt_content(str(tmp_path), "test.txt")
        assert result == test_content

    def test_extract_txt_content_with_unicode_chars(self, tmp_path):
        test_content = "Unicode content: àáâãäåæç"
        test_file = tmp_path / "unicode.txt"
        test_file.write_text(test_content, encoding="utf-8")
        files_feature = FilesFeature()
        result = files_feature._extract_txt_content(str(tmp_path), "unicode.txt")
        assert result == test_content

    def test_extract_txt_content_raises_file_not_found(self):
        files_feature = FilesFeature()
        with pytest.raises(FileNotFoundError):
            files_feature._extract_txt_content(
                "/non_existent_dir", "non_existent_file.txt"
            )

    def test_extract_docx_content_reads_file_correctly(self, tmp_path, monkeypatch):
        mock_doc = MagicMock()
        mock_paragraph1 = MagicMock()
        mock_paragraph1.text = "Paragraph 1"
        mock_paragraph2 = MagicMock()
        mock_paragraph2.text = "Paragraph 2"
        mock_doc.paragraphs = [mock_paragraph1, mock_paragraph2]

        monkeypatch.setattr(docx, "Document", lambda _: mock_doc)
        files_feature = FilesFeature()
        result = files_feature._extract_docx_content(str(tmp_path), "test.docx")

        assert result == "Paragraph 1\nParagraph 2"

    def test_extract_docx_content_skips_empty_paragraphs(self, tmp_path, monkeypatch):
        mock_doc = MagicMock()
        mock_paragraph1 = MagicMock()
        mock_paragraph1.text = "Paragraph 1"
        mock_paragraph2 = MagicMock()
        mock_paragraph2.text = "   "
        mock_paragraph3 = MagicMock()
        mock_paragraph3.text = "Paragraph 3"
        mock_doc.paragraphs = [mock_paragraph1, mock_paragraph2, mock_paragraph3]

        monkeypatch.setattr(docx, "Document", lambda _: mock_doc)
        files_feature = FilesFeature()
        result = files_feature._extract_docx_content(str(tmp_path), "test.docx")

        assert result == "Paragraph 1\nParagraph 3"

    def test_extract_docx_content_exception_handling(self, tmp_path, monkeypatch):
        def mock_document(_):
            raise Exception("Failed to open document")

        monkeypatch.setattr(docx, "Document", mock_document)

        files_feature = FilesFeature()
        with pytest.raises(Exception):
            files_feature._extract_docx_content(str(tmp_path), "invalid.docx")

    def test_extract_docx_content_with_no_paragraphs(self, tmp_path, monkeypatch):
        mock_doc = MagicMock()
        mock_doc.paragraphs = []
        monkeypatch.setattr(docx, "Document", lambda _: mock_doc)
        files_feature = FilesFeature()
        result = files_feature._extract_docx_content(str(tmp_path), "empty.docx")
        assert result == ""

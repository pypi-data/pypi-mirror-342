"""Microsoft Word (.docx) file loader for SemanticQAGen."""

import os
from typing import Dict, Any, Optional, List, Union, Tuple

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from semantic_qa_gen.document.loaders.base import BaseLoader
from semantic_qa_gen.document.models import Document, DocumentType, DocumentMetadata, Section, SectionType
from semantic_qa_gen.utils.error import DocumentError, with_error_handling


class DocxLoader(BaseLoader):
    """
    Loader for Microsoft Word (.docx) files.
    
    This loader extracts text content and metadata from Word documents,
    preserving document structure like headings and paragraphs.
    """
    
    def __init__(self, config: Optional[Dict[str, Any]] = None):
        """
        Initialize the DOCX loader.
        
        Args:
            config: Configuration dictionary for the loader.
            
        Raises:
            DocumentError: If the python-docx package is not installed.
        """
        super().__init__(config)
        
        if not DOCX_AVAILABLE:
            raise DocumentError(
                "python-docx package is required for loading .docx files.\n"
                "Install with: pip install python-docx"
            )
    
    @with_error_handling(error_types=Exception, max_retries=1)
    def load(self, path: str) -> Document:
        """
        Load a document from a DOCX file.
        
        Args:
            path: Path to the DOCX file.
            
        Returns:
            Loaded Document object.
            
        Raises:
            DocumentError: If the DOCX file cannot be loaded.
        """
        if not self.supports_type(path):
            raise DocumentError(f"Unsupported file type for DocxLoader: {path}")
        
        try:
            # Open the DOCX document
            doc = docx.Document(path)
            
            # Extract metadata
            metadata = self._extract_metadata(doc, path)
            
            # Extract content with structure preservation
            content, sections = self._extract_content_with_structure(doc)
            
            # Create the document object
            document = Document(
                content=content,
                doc_type=DocumentType.DOCX,
                path=path,
                metadata=metadata
            )
            
            # Store sections for later use in chunking
            document.sections = sections
            
            return document
            
        except Exception as e:
            raise DocumentError(f"Failed to load DOCX file {path}: {str(e)}")
    
    def supports_type(self, file_path: str) -> bool:
        """
        Check if this loader supports the given file type.
        
        Args:
            file_path: Path to the file to check.
            
        Returns:
            True if this is a DOCX file, False otherwise.
        """
        _, ext = os.path.splitext(file_path.lower())
        return ext in ['.docx']
    
    def _extract_metadata(self, doc, path: str) -> DocumentMetadata:
        """
        Extract metadata from a DOCX document.
        
        Args:
            doc: Open DOCX document.
            path: Path to the DOCX file.
            
        Returns:
            DocumentMetadata object.
        """
        # Extract document properties
        properties = doc.core_properties
        
        # Get the title, falling back to filename if not available
        title = properties.title
        if not title or title.strip() == "":
            # Fallback to filename
            file_name = os.path.basename(path)
            title, _ = os.path.splitext(file_name)
            title = title.replace('_', ' ').replace('-', ' ').strip()
            
            # Try to properly capitalize title
            if title.isupper() or title.islower():
                title = ' '.join(word.capitalize() for word in title.split())
                
        metadata = DocumentMetadata(
            title=title,
            author=properties.author,
            date=properties.created.isoformat() if properties.created else None,
            source=path,
            custom={
                'last_modified': properties.modified.isoformat() if properties.modified else None,
                'category': properties.category,
                'comments': properties.comments,
                'keywords': properties.keywords,
                'subject': properties.subject,
            }
        )
        
        return metadata
    
    def _extract_content_with_structure(self, doc) -> Tuple[str, List[Section]]:
        """
        Extract content from DOCX with structure preservation.
        
        Args:
            doc: Open DOCX document.
            
        Returns:
            Tuple of (full_text, sections_list)
        """
        # Extract paragraphs and headings
        full_text_parts = []
        sections = []
        
        heading_styles = {
            'Heading 1': 1,
            'Heading 2': 2,
            'Heading 3': 3,
            'Heading 4': 4,
            'Heading 5': 5,
            'Heading 6': 6,
            'Title': 0,  # Document title
        }
        
        for para_index, para in enumerate(doc.paragraphs):
            if not para.text.strip():
                continue
                
            # Check if this is a heading
            style_name = para.style.name if para.style else ""
            is_heading = style_name in heading_styles
            
            if is_heading:
                # This is a heading
                level = heading_styles[style_name]
                section_type = SectionType.TITLE if style_name == 'Title' else SectionType.HEADING
                
                section = Section(
                    content=para.text,
                    section_type=section_type,
                    level=level,
                    metadata={
                        'style': style_name,
                        'position': para_index
                    }
                )
                
                sections.append(section)
                full_text_parts.append(para.text)
                
            else:
                # Regular paragraph
                section = Section(
                    content=para.text,
                    section_type=SectionType.PARAGRAPH,
                    level=0,
                    metadata={
                        'style': style_name,
                        'position': para_index
                    }
                )
                
                sections.append(section)
                full_text_parts.append(para.text)
        
        # Join the text parts
        full_text = "\n\n".join(full_text_parts)
        
        return full_text, sections

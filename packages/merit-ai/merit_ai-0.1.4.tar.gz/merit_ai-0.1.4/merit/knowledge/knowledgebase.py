"""
Merit Knowledge Base

This module provides the knowledge base implementation for the Merit system.
"""
import os
import json
import csv
import sqlite3
import numpy as np
import pandas as pd
from typing import Dict, Any, List, Optional, Union, Tuple, Sequence

from ..api.base import BaseAPIClient
from ..core.models import Document
from .prompts import TOPIC_GENERATION_PROMPT
from ..core.utils import detect_language, cosine_similarity, batch_iterator
from ..core.logging import get_logger

logger = get_logger(__name__)

# Constants
DEFAULT_BATCH_SIZE = 32
DEFAULT_MIN_TOPIC_SIZE = 3
DEFAULT_LANGUAGE_DETECTION_SAMPLE_SIZE = 10
DEFAULT_LANGUAGE_DETECTION_MAX_TEXT_LENGTH = 300

# Try to import optional dependencies
try:
    import pymongo
    MONGODB_AVAILABLE = True
except ImportError:
    MONGODB_AVAILABLE = False

try:
    from sklearn.feature_extraction.text import TfidfVectorizer
    TFIDF_AVAILABLE = True
except ImportError:
    TFIDF_AVAILABLE = False

class KnowledgeBase:
    """
    A knowledge base for the Merit RAG system.
    
    This class provides functionality for creating and managing a knowledge base,
    including embedding documents, finding topics, and searching for relevant documents.
    """
    
    @classmethod
    def from_knowledge_bases(cls, knowledge_bases: List["KnowledgeBase"], client: BaseAPIClient = None) -> "KnowledgeBase":
        """
        Create a combined knowledge base from a list of knowledge bases.
        
        Args:
            knowledge_bases: A list of knowledge bases to combine.
            client: The API client to use for the combined knowledge base. If None, uses the client from the first knowledge base.
            
        Returns:
            KnowledgeBase: A new knowledge base containing all documents from the input knowledge bases.
            
        Raises:
            ValueError: If the list of knowledge bases is empty.
        """
        if not knowledge_bases:
            raise ValueError("Cannot create a knowledge base from an empty list of knowledge bases")
        
        # Use the client from the first knowledge base if not provided
        if client is None:
            client = knowledge_bases[0]._client
        
        # Collect all documents from all knowledge bases
        all_documents = []
        for kb in knowledge_bases:
            all_documents.extend(kb.documents)
        
        # Create document dictionaries
        document_dicts = []
        for doc in all_documents:
            document_dict = {
                "content": doc.content,
                "metadata": doc.metadata,
                "id": doc.id
            }
            document_dicts.append(document_dict)
        
        # Create a new knowledge base with all documents
        return cls(data=document_dicts, client=client)
    
    def __init__(
        self,
        data: Union[pd.DataFrame, List[Dict[str, Any]], str, Dict[str, Any]],
        client: BaseAPIClient,
        columns: Optional[Sequence[str]] = None,
        seed: Optional[int] = None,
        min_topic_size: Optional[int] = None,
        batch_size: int = DEFAULT_BATCH_SIZE,
        # New parameters for data source configuration
        source_type: Optional[str] = None,  # 'csv', 'json', 'sqlite', 'mongodb', 'file'
        source_config: Optional[Dict[str, Any]] = None,
    ):
        """
        Initialize the knowledge base.
        
        Args:
            data: The data to create the knowledge base from. Can be:
                - pandas DataFrame
                - List of dictionaries
                - Path to a file (CSV, JSON, PDF, TXT, DOCX)
                - SQLite connection string or config dict
                - MongoDB connection config dict
            client: The API client to use for embeddings and text generation.
            columns: The columns to use from the data. If None, all columns are used.
            seed: The random seed to use.
            min_topic_size: The minimum number of documents to form a topic.
            batch_size: The batch size to use for embeddings.
            source_type: Explicitly specify the source type. If None, will be auto-detected.
            source_config: Additional configuration for the data source.
        """
        # Process data based on type
        source_config = source_config or {}
        
        if isinstance(data, str):
            # Handle file paths or connection strings
            if source_type is None:
                # Auto-detect source type from file extension
                ext = os.path.splitext(data)[1].lower()
                if ext == '.csv':
                    source_type = 'csv'
                elif ext == '.json':
                    source_type = 'json'
                elif ext in ['.db', '.sqlite', '.sqlite3']:
                    source_type = 'sqlite'
                elif 'mongodb://' in data or 'mongodb+srv://' in data:
                    source_type = 'mongodb'
                else:
                    # Default to treating as a file
                    source_type = 'file'
            
            # Load data from the specified source
            if source_type == 'csv':
                data = self._load_from_csv(data, source_config)
            elif source_type == 'json':
                data = self._load_from_json(data, source_config)
            elif source_type == 'sqlite':
                data = self._load_from_sqlite(data, source_config)
            elif source_type == 'mongodb':
                if not MONGODB_AVAILABLE:
                    raise ImportError("MongoDB support requires pymongo. Install with 'pip install pymongo'")
                data = self._load_from_mongodb(data, source_config)
            elif source_type == 'file':
                data = self._load_from_file(data, source_config)
        
        elif isinstance(data, dict) and source_type == 'mongodb':
            # Handle MongoDB config dictionary
            if not MONGODB_AVAILABLE:
                raise ImportError("MongoDB support requires pymongo. Install with 'pip install pymongo'")
            data = self._load_from_mongodb(None, data)
        
        # Convert data to DataFrame if it's a list
        if isinstance(data, list):
            data = pd.DataFrame(data)
        
        if len(data) == 0:
            raise ValueError("Cannot create a knowledge base from empty data")
        
        # Set up random number generator
        self._rng = np.random.default_rng(seed=seed)
        
        # Store parameters
        self._client = client
        self._batch_size = batch_size
        self._min_topic_size = min_topic_size or DEFAULT_MIN_TOPIC_SIZE
        
        # Create documents
        self._documents = self._create_documents(data, columns)
        
        if len(self._documents) == 0:
            raise ValueError("Cannot create a knowledge base with empty documents")
        
        # Create document index
        self._document_index = {doc.id: doc for doc in self._documents}
        
        # Initialize caches
        self._embeddings_cache = None
        self._topics_cache = None
        self._index_cache = None
        self._reduced_embeddings_cache = None
        self._tfidf_vectorizer = None
        self._tfidf_matrix = None
        
        # Detect language
        self._language = self._detect_language()
        
        logger.info(f"Created knowledge base with {len(self._documents)} documents in language '{self._language}'")
    
    def _load_from_csv(self, file_path: str, config: Dict[str, Any]) -> pd.DataFrame:
        """
        Load data from a CSV file.
        
        Args:
            file_path: Path to the CSV file
            config: Configuration for loading CSV
                - content_columns: Columns to use as document content
                - metadata_columns: Columns to include in document metadata
                - id_column: Column to use as document ID
                - Any other pandas.read_csv parameters
                
        Returns:
            pd.DataFrame: Loaded data
        """
        # Extract CSV-specific config
        content_columns = config.get('content_columns')
        metadata_columns = config.get('metadata_columns')
        id_column = config.get('id_column')
        
        # Remove our custom params from kwargs to pass to pandas
        csv_kwargs = {k: v for k, v in config.items() 
                     if k not in ['content_columns', 'metadata_columns', 'id_column']}
        
        # Load CSV
        df = pd.read_csv(file_path, **csv_kwargs)
        
        # Process the DataFrame to prepare it for the knowledge base
        return self._process_dataframe(df, content_columns, metadata_columns, id_column)
    
    def _load_from_json(self, file_path: str, config: Dict[str, Any]) -> pd.DataFrame:
        """
        Load data from a JSON file.
        
        Args:
            file_path: Path to the JSON file
            config: Configuration for loading JSON
                - content_fields: Fields to use as document content
                - metadata_fields: Fields to include in document metadata
                - id_field: Field to use as document ID
                - records_path: Path to the records in the JSON (e.g., 'data.records')
                
        Returns:
            pd.DataFrame: Loaded data
        """
        # Extract JSON-specific config
        content_fields = config.get('content_fields')
        metadata_fields = config.get('metadata_fields')
        id_field = config.get('id_field')
        records_path = config.get('records_path')
        
        # Load JSON
        with open(file_path, 'r', encoding='utf-8') as f:
            json_data = json.load(f)
        
        # Navigate to records if a path is specified
        if records_path:
            parts = records_path.split('.')
            for part in parts:
                if part in json_data:
                    json_data = json_data[part]
                else:
                    raise ValueError(f"Path '{records_path}' not found in JSON")
        
        # Ensure we have a list of records
        if not isinstance(json_data, list):
            if isinstance(json_data, dict):
                # Try to convert a single record to a list
                json_data = [json_data]
            else:
                raise ValueError("JSON data must be a list of records or a single record")
        
        # Convert to DataFrame
        df = pd.DataFrame(json_data)
        
        # Process the DataFrame to prepare it for the knowledge base
        return self._process_dataframe(df, content_fields, metadata_fields, id_field)
    
    def _load_from_sqlite(self, db_path: str, config: Dict[str, Any]) -> pd.DataFrame:
        """
        Load data from a SQLite database.
        
        Args:
            db_path: Path to the SQLite database
            config: Configuration for loading SQLite
                - query: SQL query to execute
                - content_columns: Columns to use as document content
                - metadata_columns: Columns to include in document metadata
                - id_column: Column to use as document ID
                
        Returns:
            pd.DataFrame: Loaded data
        """
        # Extract SQLite-specific config
        query = config.get('query')
        content_columns = config.get('content_columns')
        metadata_columns = config.get('metadata_columns')
        id_column = config.get('id_column')
        
        if not query:
            raise ValueError("SQLite source requires a 'query' parameter")
        
        # Connect to database
        conn = sqlite3.connect(db_path)
        
        try:
            # Execute query
            df = pd.read_sql_query(query, conn)
            
            # Process the DataFrame to prepare it for the knowledge base
            return self._process_dataframe(df, content_columns, metadata_columns, id_column)
        finally:
            conn.close()
    
    def _load_from_mongodb(self, connection_string: Optional[str], config: Dict[str, Any]) -> pd.DataFrame:
        """
        Load data from a MongoDB collection.
        
        Args:
            connection_string: MongoDB connection string (can be None if provided in config)
            config: Configuration for loading MongoDB
                - connection_string: MongoDB connection string (if not provided as first arg)
                - database: Database name
                - collection: Collection name
                - query: MongoDB query to filter documents
                - content_fields: Fields to use as document content
                - metadata_fields: Fields to include in document metadata
                - id_field: Field to use as document ID
                
        Returns:
            pd.DataFrame: Loaded data
        """
        # Extract MongoDB-specific config
        conn_str = connection_string or config.get('connection_string')
        database = config.get('database')
        collection = config.get('collection')
        query = config.get('query', {})
        content_fields = config.get('content_fields')
        metadata_fields = config.get('metadata_fields')
        id_field = config.get('id_field', '_id')
        
        if not conn_str:
            raise ValueError("MongoDB source requires a connection string")
        if not database:
            raise ValueError("MongoDB source requires a 'database' parameter")
        if not collection:
            raise ValueError("MongoDB source requires a 'collection' parameter")
        
        # Connect to MongoDB
        client = pymongo.MongoClient(conn_str)
        
        try:
            # Get database and collection
            db = client[database]
            coll = db[collection]
            
            # Execute query
            documents = list(coll.find(query))
            
            # Convert ObjectId to string for JSON serialization
            for doc in documents:
                if '_id' in doc and not isinstance(doc['_id'], str):
                    doc['_id'] = str(doc['_id'])
            
            # Convert to DataFrame
            df = pd.DataFrame(documents)
            
            # Process the DataFrame to prepare it for the knowledge base
            return self._process_dataframe(df, content_fields, metadata_fields, id_field)
        finally:
            client.close()
    
    def _load_from_file(self, file_path: str, config: Dict[str, Any]) -> pd.DataFrame:
        """
        Load data from a file (PDF, TXT, DOCX).
        
        Args:
            file_path: Path to the file
            config: Configuration for loading the file
                - chunk_size: Size of text chunks to create (default: 1000)
                - chunk_overlap: Overlap between chunks (default: 200)
                
        Returns:
            pd.DataFrame: Loaded data
        """
        # Extract file-specific config
        chunk_size = config.get('chunk_size', 1000)
        chunk_overlap = config.get('chunk_overlap', 200)
        
        # Get file extension
        ext = os.path.splitext(file_path)[1].lower()
        
        # Read file content based on extension
        if ext == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
        elif ext == '.pdf':
            try:
                import PyPDF2
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text() + "\n\n"
            except ImportError:
                raise ImportError("PDF support requires PyPDF2. Install with 'pip install PyPDF2'")
        elif ext == '.docx':
            try:
                import docx
                doc = docx.Document(file_path)
                text = "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            except ImportError:
                raise ImportError("DOCX support requires python-docx. Install with 'pip install python-docx'")
        else:
            raise ValueError(f"Unsupported file type: {ext}")
        
        # Chunk the text
        chunks = self._chunk_text(text, chunk_size, chunk_overlap)
        
        # Create a DataFrame with the chunks
        data = [{"content": chunk, "source": file_path, "chunk_index": i} for i, chunk in enumerate(chunks)]
        
        return pd.DataFrame(data)
    
    def _chunk_text(self, text: str, chunk_size: int, chunk_overlap: int) -> List[str]:
        """
        Split text into overlapping chunks.
        
        Args:
            text: The text to split
            chunk_size: Size of each chunk
            chunk_overlap: Overlap between chunks
            
        Returns:
            List[str]: List of text chunks
        """
        if not text:
            return []
            
        # Split text into chunks
        chunks = []
        start = 0
        
        while start < len(text):
            # Get chunk
            end = min(start + chunk_size, len(text))
            chunk = text[start:end]
            
            # Add chunk to list
            if chunk.strip():
                chunks.append(chunk)
            
            # Move to next chunk
            start = end - chunk_overlap
            
            # Avoid getting stuck in an infinite loop
            if start >= end:
                break
        
        return chunks
    
    def _process_dataframe(self, df: pd.DataFrame, 
                          content_columns: Optional[List[str]] = None,
                          metadata_columns: Optional[List[str]] = None,
                          id_column: Optional[str] = None) -> pd.DataFrame:
        """
        Process a DataFrame to prepare it for the knowledge base.
        
        Args:
            df: Input DataFrame
            content_columns: Columns to use as document content
            metadata_columns: Columns to include in document metadata
            id_column: Column to use as document ID
            
        Returns:
            pd.DataFrame: Processed DataFrame ready for knowledge base
        """
        # Handle empty DataFrame
        if df.empty:
            return df
        
        # If content_columns is not specified, use all columns except id_column
        if content_columns is None:
            if id_column:
                content_columns = [col for col in df.columns if col != id_column]
            else:
                content_columns = df.columns.tolist()
        
        # If metadata_columns is not specified, use all columns
        if metadata_columns is None:
            metadata_columns = df.columns.tolist()
        
        # Create a new DataFrame with the processed data
        processed_data = []
        
        for idx, row in df.iterrows():
            # Create content by joining content columns
            if len(content_columns) > 1:
                content = "\n".join(f"{col}: {row[col]}" for col in content_columns if col in row)
            else:
                content = str(row[content_columns[0]])
            
            # Skip empty documents
            if not content.strip():
                continue
            
            # Create metadata
            metadata = {col: row[col] for col in metadata_columns if col in row}
            
            # Get ID - use the specified id_column if available
            if id_column and id_column in row and not pd.isna(row[id_column]):
                doc_id = str(row[id_column])
            else:
                doc_id = str(idx)
            
            # Create document dictionary
            doc_dict = {
                "content": content,
                "metadata": metadata,
                "id": doc_id
            }
            
            processed_data.append(doc_dict)
        
        return pd.DataFrame(processed_data)
    
    def _create_documents(self, data: pd.DataFrame, columns: Optional[Sequence[str]] = None) -> List[Document]:
        """
        Create documents from the data.
        
        Args:
            data: The data to create documents from.
            columns: The columns to use from the data. If None, all columns are used.
            
        Returns:
            List[Document]: The created documents.
        """
        # If columns is None, use all columns
        if columns is None:
            columns = data.columns.tolist()
        
        # Create documents
        documents = []
        for idx, row in data.iterrows():
            # Create content by joining columns
            if len(columns) > 1:
                content = "\n".join(f"{col}: {row[col]}" for col in columns if col in row)
            else:
                content = str(row[columns[0]])
            
            # Skip empty documents
            if not content.strip():
                continue
            
            # Get ID from the row if available
            if 'id' in row and not pd.isna(row['id']):
                doc_id = str(row['id'])
            else:
                # Check if the document ID is in the metadata
                if 'metadata' in row and isinstance(row['metadata'], dict) and 'document_id' in row['metadata']:
                    doc_id = str(row['metadata']['document_id'])
                else:
                    doc_id = str(idx)
            
            # Create document
            doc = Document(
                content=content,
                metadata=row.to_dict(),
                id=doc_id,
            )
            documents.append(doc)
        
        return documents
    
    def _detect_language(self) -> str:
        """
        Detect the language of the documents.
        
        Returns:
            str: The detected language code (e.g., "en", "fr").
        """
        # Sample documents for language detection
        sample_size = min(DEFAULT_LANGUAGE_DETECTION_SAMPLE_SIZE, len(self._documents))
        sample_docs = self._rng.choice(self._documents, size=sample_size, replace=False)
        
        # Detect language for each document
        languages = []
        for doc in sample_docs:
            # Use only the first N characters for faster detection
            text = doc.content[:DEFAULT_LANGUAGE_DETECTION_MAX_TEXT_LENGTH]
            lang = detect_language(text)
            languages.append(lang)
        
        # Count language occurrences
        lang_counts = {}
        for lang in languages:
            lang_counts[lang] = lang_counts.get(lang, 0) + 1
        
        # Return the most common language, or "en" if no language is detected
        if not lang_counts:
            return "en"
        
        return max(lang_counts.items(), key=lambda x: x[1])[0]
    
    @property
    def documents(self) -> List[Document]:
        """
        Get all documents in the knowledge base, as a list of Document objects.
        Usage: 
        kb = KnowledgeBase(data, client)
        for doc in kb.documents:
            print(doc.content)
        Returns:
            List[Document]: All documents in the knowledge base.
        """
        return self._documents
    
    def get_all_documents(self) -> List[Document]:
        """
        Get all documents in the knowledge base.
        
        This method provides an alternative to the documents property.
        
        Returns:
            List[Document]: All documents in the knowledge base.
        """
        return self._documents
    
    @property
    def language(self) -> str:
        """
        Get the language of the knowledge base.
        
        Returns:
            str: The language code (e.g., "en", "fr").
        """
        return self._language
    
    @property
    def embeddings(self) -> np.ndarray:
        """
        Get the embeddings of all documents in the knowledge base.
        
        Returns:
            np.ndarray: The embeddings of all documents.
        """
        if self._embeddings_cache is not None:
            return self._embeddings_cache
            
        logger.info("Computing embeddings for knowledge base")
        
        # Get embeddings in batches
        all_embeddings = []
        total_batches = (len(self._documents) + self._batch_size - 1) // self._batch_size
        
        for batch_idx, batch in enumerate(batch_iterator(self._documents, self._batch_size)):
            logger.info(f"Processing batch {batch_idx+1}/{total_batches}")
            
            batch_texts = [doc.content for doc in batch]
            batch_embeddings = self._client.get_embeddings(batch_texts)
            all_embeddings.extend(batch_embeddings)
        
        # Store embeddings in documents
        for doc, emb in zip(self._documents, all_embeddings):
            doc.embeddings = emb
        
        # Cache embeddings
        self._embeddings_cache = np.array(all_embeddings)
        
        return self._embeddings_cache
        
    def set_embeddings(self, embeddings: List[List[float]]) -> None:
        """
        Set pre-computed embeddings for the knowledge base.
        
        This method allows you to provide pre-computed embeddings instead of
        generating them on the fly, which can be useful for large knowledge bases
        or when you want to reuse embeddings across multiple sessions.
        
        Args:
            embeddings: List of embeddings, one for each document in the knowledge base.
                The order must match the order of documents in the knowledge base.
                
        Raises:
            ValueError: If the number of embeddings doesn't match the number of documents.
        """
        if len(embeddings) != len(self._documents):
            raise ValueError(f"Number of embeddings ({len(embeddings)}) doesn't match number of documents ({len(self._documents)})")
            
        # Store embeddings in documents
        for doc, emb in zip(self._documents, embeddings):
            doc.embeddings = emb
            
        # Cache embeddings
        self._embeddings_cache = np.array(embeddings)
        
        logger.info(f"Set {len(embeddings)} pre-computed embeddings for knowledge base")
    
    def get_reduced_embeddings(
        self, 
        n_components: int = 2, 
        method: str = "pca", 
        random_state: int = 42,
        force_recompute: bool = False
    ) -> np.ndarray:
        """
        Get the reduced embeddings of all documents in the knowledge base.
        
        This method reduces the dimensionality of document embeddings for visualization
        or clustering purposes. It supports multiple dimensionality reduction methods.
        
        Args:
            n_components: Number of dimensions to reduce to (default: 2)
            method: Dimensionality reduction method to use:
                - "pca": Principal Component Analysis (default)
                - "tsne": t-SNE (t-Distributed Stochastic Neighbor Embedding)
                - "umap": Uniform Manifold Approximation and Projection
            random_state: Random seed for reproducibility
            force_recompute: Whether to force recomputation even if cached
            
        Returns:
            np.ndarray: The reduced embeddings of all documents
            
        Raises:
            ImportError: If the requested method is not available
            ValueError: If the method is not supported
        """
        # Return cached embeddings if available and not forcing recomputation
        if self._reduced_embeddings_cache is not None and not force_recompute:
            # Check if the cached embeddings have the right number of components
            if self._reduced_embeddings_cache.shape[1] == n_components:
                return self._reduced_embeddings_cache
        
        logger.info(f"Computing reduced embeddings using {method} with {n_components} components")
        
        try:
            # Get embeddings
            embeddings = self.embeddings
            
            # Check if we have enough documents for dimensionality reduction
            if len(embeddings) < 2:
                logger.warning("Not enough documents for dimensionality reduction")
                self._reduced_embeddings_cache = np.zeros((len(self._documents), n_components))
                return self._reduced_embeddings_cache
            
            # Clean data before dimensionality reduction
            embeddings_clean = np.nan_to_num(embeddings, nan=0.0, posinf=1e10, neginf=-1e10)
            
            # Apply dimensionality reduction based on method
            if method == "pca":
                # Use PCA for dimensionality reduction
                from sklearn.decomposition import PCA
                reducer = PCA(n_components=n_components, random_state=random_state)
                reduced = reducer.fit_transform(embeddings_clean)
                
            elif method == "tsne":
                # Use t-SNE for dimensionality reduction
                try:
                    from sklearn.manifold import TSNE
                    reducer = TSNE(n_components=n_components, random_state=random_state)
                    reduced = reducer.fit_transform(embeddings_clean)
                except ImportError:
                    logger.warning("t-SNE not available, falling back to PCA")
                    from sklearn.decomposition import PCA
                    reducer = PCA(n_components=n_components, random_state=random_state)
                    reduced = reducer.fit_transform(embeddings_clean)
                    
            elif method == "umap":
                # Use UMAP for dimensionality reduction
                try:
                    import umap
                    reducer = umap.UMAP(n_components=n_components, random_state=random_state)
                    reduced = reducer.fit_transform(embeddings_clean)
                except ImportError:
                    logger.warning("UMAP not available, falling back to PCA")
                    from sklearn.decomposition import PCA
                    reducer = PCA(n_components=n_components, random_state=random_state)
                    reduced = reducer.fit_transform(embeddings_clean)
            else:
                raise ValueError(f"Unsupported dimensionality reduction method: {method}")
            
            # Store reduced embeddings in documents
            for doc, emb in zip(self._documents, reduced):
                doc.reduced_embeddings = emb.tolist()
            
            # Cache reduced embeddings
            self._reduced_embeddings_cache = reduced
            
        except Exception as e:
            logger.error(f"Failed to compute reduced embeddings: {str(e)}")
            # Return empty array as fallback
            self._reduced_embeddings_cache = np.zeros((len(self._documents), n_components))
        
        return self._reduced_embeddings_cache
        
    @property
    def reduced_embeddings(self) -> np.ndarray:
        """
        Get the reduced embeddings of all documents in the knowledge base.
        
        This is a convenience property that calls get_reduced_embeddings with default parameters.
        For more control over the dimensionality reduction process, use get_reduced_embeddings directly.
        
        Returns:
            np.ndarray: The reduced embeddings of all documents (2D by default).
        """
        return self.get_reduced_embeddings(n_components=2, method="pca")
    
    @property
    def topics(self) -> Dict[int, str]:
        """
        Get the topics of the knowledge base.
        
        Returns:
            Dict[int, str]: A dictionary mapping topic IDs to topic names.
        """
        if self._topics_cache is None:
            logger.info("Finding topics in knowledge base")
            self._topics_cache = self._find_topics()
        
        return self._topics_cache
    def _find_topics(self) -> Dict[int, str]:
        """
        Find topics in the knowledge base.
        
        Returns:
            Dict[int, str]: A dictionary mapping topic IDs to topic names.
        """
        try:
            # Try to import hdbscan
            try:
                from hdbscan import HDBSCAN
                hdbscan_available = True
            except ImportError:
                logger.warning("hdbscan not installed, falling back to KMeans clustering")
                hdbscan_available = False
            
            import numpy as np
            
            # Get reduced embeddings
            embeddings = self.reduced_embeddings
            
            if hdbscan_available:
                # Use HDBSCAN for clustering
                try:
                    from sklearn.neighbors import NearestNeighbors
                    
                    # Create HDBSCAN clusterer with more relaxed parameters to reduce noise points
                    clusterer = HDBSCAN(
                        min_cluster_size=max(2, self._min_topic_size),
                        min_samples=2,  # Reduced from 3 to be less strict
                        metric="euclidean",
                        cluster_selection_epsilon=0.5,  # Increased from 0.0 to be more inclusive
                        cluster_selection_method='eom'  # Excess of Mass - tends to find more varied cluster sizes
                    )
                    
                    # Cluster documents
                    clustering = clusterer.fit(embeddings)
                    labels = clustering.labels_
                    
                    # Handle noise points (-1 labels) by assigning them to the nearest cluster
                    noise_indices = np.where(labels == -1)[0]
                    if len(noise_indices) > 0:
                        logger.info(f"Found {len(noise_indices)} noise points, assigning to nearest clusters")
                        
                        # Get non-noise cluster indices
                        valid_clusters = np.unique(labels[labels != -1])
                        
                        if len(valid_clusters) > 0:
                            # For each noise point, find the nearest non-noise point
                            for idx in noise_indices:
                                # Get embeddings for this noise point
                                noise_embedding = embeddings[idx].reshape(1, -1)
                                
                                # Find distances to all points
                                distances = np.linalg.norm(embeddings - noise_embedding, axis=1)
                                
                                # Sort by distance (excluding self)
                                sorted_indices = np.argsort(distances)
                                
                                # Find the nearest non-noise point
                                for nearest_idx in sorted_indices:
                                    if labels[nearest_idx] != -1 and nearest_idx != idx:
                                        # Assign this noise point to the same cluster as its nearest non-noise neighbor
                                        labels[idx] = labels[nearest_idx]
                                        break
                        else:
                            # If all points are noise, create a single cluster
                            logger.warning("All points classified as noise, creating a single cluster")
                            labels = np.zeros_like(labels)
                except Exception as e:
                    logger.warning(f"HDBSCAN clustering failed: {str(e)}, falling back to KMeans")
                    hdbscan_available = False
            
            # Fallback to KMeans if HDBSCAN is not available or failed
            if not hdbscan_available:
                try:
                    from sklearn.cluster import KMeans
                    
                    # Determine number of clusters
                    n_docs = len(self._documents)
                    n_clusters = min(max(2, n_docs // 10), 20)  # Between 2 and 20 clusters
                    
                    # Create KMeans clusterer
                    clusterer = KMeans(n_clusters=n_clusters, random_state=42)
                    
                    # Cluster documents
                    labels = clusterer.fit_predict(embeddings)
                    
                    logger.info(f"Used KMeans clustering with {n_clusters} clusters")
                except Exception as e:
                    logger.error(f"KMeans clustering failed: {str(e)}, using simple clustering")
                    # Simple clustering as last resort
                    n_docs = len(self._documents)
                    n_clusters = min(max(2, n_docs // 10), 20)
                    labels = np.array([i % n_clusters for i in range(n_docs)])
            
            # Assign topic IDs to documents
            for i, doc in enumerate(self._documents):
                doc.topic_id = int(labels[i])
            
            # Get unique topic IDs
            topic_ids = set(labels)
            
            # Generate topic names
            topics = {}
            for topic_id in topic_ids:
                if topic_id == -1:
                    # -1 is the noise cluster (should be few or none after our processing)
                    topics[topic_id] = "Other"
                else:
                    # Get documents in this topic
                    topic_docs = [doc for doc in self._documents if doc.topic_id == topic_id]
                    # Generate topic name
                    topic_name = self._generate_topic_name(topic_docs)
                    topics[topic_id] = topic_name
            
            logger.info(f"Found {len(topics)} topics in knowledge base")
            return topics
        
        except Exception as e:
            logger.error(f"Failed to find topics: {str(e)}")
            # Return a single "Unknown" topic as fallback
            for doc in self._documents:
                doc.topic_id = 0
            return {0: "Unknown"}
    
    def _generate_topic_name(self, topic_documents: List[Document]) -> str:
        """
        Generate a name for a topic.
        
        Args:
            topic_documents: The documents in the topic.
            
        Returns:
            str: The generated topic name.
        """
        # Shuffle documents to get a random sample
        self._rng.shuffle(topic_documents)
        
        # Get a sample of documents
        sample_size = min(10, len(topic_documents))
        sample_docs = topic_documents[:sample_size]
        
        # Create prompt
        topics_str = "\n\n".join(["----------" + doc.content[:500] for doc in sample_docs])
        
        # Prevent context window overflow
        topics_str = topics_str[:3 * 8192]
        
        prompt = TOPIC_GENERATION_PROMPT.safe_format(
            language=self._language,
            topics_elements=topics_str
        )
        
        try:
            # Generate topic name
            topic_name = self._client.generate_text(prompt)
            
            # Clean up topic name
            topic_name = topic_name.strip().strip('"')
            
            if not topic_name:
                logger.warning("Generated empty topic name, using fallback")
                return "Unknown Topic"
            
            logger.info(f"Generated topic name: {topic_name}")
            return topic_name
        
        except Exception as e:
            logger.error(f"Failed to generate topic name: {str(e)}")
            return "Unknown Topic"
    
    def get_document(self, doc_id: str) -> Optional[Document]:
        """
        Get a document by ID.
        
        Args:
            doc_id: The ID of the document to get.
            
        Returns:
            Optional[Document]: The document, or None if not found.
        """
        return self._document_index.get(doc_id)
    
    def get_documents_by_topic(self, topic_id: int) -> List[Document]:
        """
        Get all documents in a topic.
        
        Args:
            topic_id: The ID of the topic to get documents for.
            
        Returns:
            List[Document]: The documents in the topic.
        """
        return [doc for doc in self._documents if doc.topic_id == topic_id]
    
    def get_random_document(self) -> Document:
        """
        Get a random document from the knowledge base.
        
        Returns:
            Document: A random document.
        """
        return self._rng.choice(self._documents)
    
    def get_random_documents(self, n: int, with_replacement: bool = False) -> List[Document]:
        """
        Get random documents from the knowledge base.
        
        Args:
            n: The number of documents to get.
            with_replacement: Whether to allow the same document to be selected multiple times.
            
        Returns:
            List[Document]: The random documents.
        """
        if with_replacement or n > len(self._documents):
            return list(self._rng.choice(self._documents, n, replace=True))
        else:
            return list(self._rng.choice(self._documents, n, replace=False))
    def get_representative_documents(self, n: int, exclude_noise: bool = True) -> List[Document]:
        """
        Get a representative sample of documents based on topic distribution.
        
        This method samples documents from each topic proportionally to the
        number of documents in that topic, ensuring the sample represents
        the knowledge base structure.
        
        Args:
            n: The number of documents to sample.
            exclude_noise: Whether to exclude documents from the noise topic (topic_id=-1).
            
        Returns:
            List[Document]: The sampled documents.
        """
        # Use the new sample_documents_with_input_allocation method and extract just the documents
        doc_input_pairs = self.sample_documents_with_input_allocation(
            n=n,
            strategy="representative"
        )
        return [doc for doc, _ in doc_input_pairs]
    
    def sample_documents_with_input_allocation(
        self, 
        n: int, 
        strategy: str,
        items_per_document: int = 1
    ) -> List[Tuple[Document, int]]:
        """
        Sample documents from the knowledge base and determine input allocation.
        
        This method samples documents according to the specified strategy and
        determines how many test inputs should be generated for each document.
        
        Args:
            n: The total number of test inputs to generate.
            strategy: The sampling strategy to use:
                - "random": Randomly sample documents.
                - "representative": Sample documents proportionally to topic distribution.
                - "per_document": Sample documents and generate a fixed number of inputs per document.
            items_per_document: Number of items to generate per document.
                Only used when strategy is "per_document".
                
        Returns:
            List[Tuple[Document, int]]: List of (document, inputs_to_generate) tuples.
        """
        if strategy == "random":
            return self._sample_random_with_allocation(n)
        elif strategy == "representative":
            return self._sample_representative_with_allocation(n)
        elif strategy == "per_document":
            return self._sample_per_document_with_allocation(n, items_per_document)
        else:
            raise ValueError(f"Unknown sampling strategy: {strategy}")
    
    def _sample_random_with_allocation(self, n: int) -> List[Tuple[Document, int]]:
        """
        Randomly sample documents and determine input allocation.
        
        Args:
            n: The total number of test inputs to generate.
            
        Returns:
            List[Tuple[Document, int]]: List of (document, inputs_to_generate) tuples.
        """
        # Determine how many documents to sample
        docs_to_sample = min(n, len(self._documents))
        
        # Sample documents
        sampled_docs = self.get_random_documents(docs_to_sample)
        
        # Calculate inputs per document
        base_inputs_per_doc = n // docs_to_sample
        extra_inputs = n % docs_to_sample
        
        # Allocate inputs to documents
        result = []
        for i, doc in enumerate(sampled_docs):
            inputs_for_doc = base_inputs_per_doc
            if i < extra_inputs:
                inputs_for_doc += 1
            result.append((doc, inputs_for_doc))
        
        return result

    def _sample_representative_with_allocation(self, n: int) -> List[Tuple[Document, int]]:
        """
        Sample documents proportionally to topic distribution and determine input allocation.
        
        Args:
            n: The total number of test inputs to generate.
            
        Returns:
            List[Tuple[Document, int]]: List of (document, inputs_to_generate) tuples.
        """
        # Get all topics
        topics = self.topics
        
        # Filter out noise topic
        valid_topics = {topic_id: name for topic_id, name in topics.items() if topic_id != -1}
        
        if not valid_topics:
            logger.warning("No valid topics found, falling back to random sampling")
            return self._sample_random_with_allocation(n)
        
        # Calculate the total number of documents in valid topics
        topic_document_counts = {}
        total_documents = 0
        
        for topic_id in valid_topics:
            topic_docs = self.get_documents_by_topic(topic_id)
            topic_document_counts[topic_id] = len(topic_docs)
            total_documents += len(topic_docs)
        
        # Calculate how many inputs to generate for each topic
        # based on the proportion of documents in that topic
        topic_input_counts = {}
        remaining_inputs = n
        
        # First, ensure each topic gets at least one input
        for topic_id in valid_topics:
            topic_input_counts[topic_id] = 1
            remaining_inputs -= 1
        
        # If we've already allocated all inputs, we're done
        if remaining_inputs <= 0:
            remaining_inputs = 0
        else:
            # Distribute remaining inputs proportionally
            for topic_id, doc_count in topic_document_counts.items():
                proportion = doc_count / total_documents
                additional_inputs = int(proportion * remaining_inputs)
                topic_input_counts[topic_id] += additional_inputs
                remaining_inputs -= additional_inputs
        
        # If we have any remaining inputs due to rounding,
        # distribute them to the largest topics
        if remaining_inputs > 0:
            sorted_topics = sorted(
                topic_document_counts.items(),
                key=lambda x: x[1],
                reverse=True
            )
            
            for topic_id, _ in sorted_topics:
                if remaining_inputs <= 0:
                    break
                topic_input_counts[topic_id] += 1
                remaining_inputs -= 1
        
        # Now determine how many documents to sample from each topic
        # and how many inputs to generate per document
        result = []
        
        for topic_id, input_count in topic_input_counts.items():
            topic_docs = self.get_documents_by_topic(topic_id)
            
            # Determine how many documents to sample from this topic
            # We want to maximize document diversity, so we'll use as many
            # documents as possible (up to the number of inputs needed)
            docs_to_sample = min(len(topic_docs), input_count)
            
            # Sample documents from this topic
            sampled_indices = self._rng.choice(
                len(topic_docs),
                size=docs_to_sample,
                replace=False
            )
            sampled_docs = [topic_docs[i] for i in sampled_indices]
            
            # Calculate inputs per document
            base_inputs_per_doc = input_count // docs_to_sample
            extra_inputs = input_count % docs_to_sample
            
            # Distribute inputs across documents
            for i, doc in enumerate(sampled_docs):
                # Give extra inputs to the first 'extra_inputs' documents
                inputs_for_this_doc = base_inputs_per_doc
                if i < extra_inputs:
                    inputs_for_this_doc += 1
                    
                result.append((doc, inputs_for_this_doc))
        
        # Shuffle the result to avoid clustering by topic
        self._rng.shuffle(result)
        
        return result

    def _sample_per_document_with_allocation(self, n: int, items_per_document: int) -> List[Tuple[Document, int]]:
        """
        Sample documents and allocate a fixed number of inputs per document.
        
        Args:
            n: The total number of test inputs to generate.
            items_per_document: Number of items to generate per document.
            
        Returns:
            List[Tuple[Document, int]]: List of (document, inputs_to_generate) tuples.
        """
        # Calculate how many documents to sample
        docs_to_sample = (n + items_per_document - 1) // items_per_document
        docs_to_sample = min(docs_to_sample, len(self._documents))
        
        # Sample documents
        sampled_docs = self.get_random_documents(docs_to_sample)
        
        # Allocate inputs to documents
        result = []
        remaining_inputs = n
        
        for doc in sampled_docs:
            # Allocate items to this document
            items_for_doc = min(items_per_document, remaining_inputs)
            if items_for_doc > 0:
                result.append((doc, items_for_doc))
                remaining_inputs -= items_for_doc
            
            # Stop if we've allocated all inputs
            if remaining_inputs <= 0:
                break
        
        return result
    def search(self, query, k: int = 5, mode: str = "embedding") -> List[Tuple[Document, float]]:
        """
        Search for documents similar to the query.
        
        Args:
            query: The query to search for. Can be a string or an Input object.
            k: The number of results to return.
            mode: The search mode to use. Options:
                - "embedding": Use embeddings for semantic search (default)
                - "keyword": Use keyword matching
                - "hybrid": Use both embedding and keyword search
            
        Returns:
            List[Tuple[Document, float]]: The search results, as (document, score) pairs.
            
        Raises:
            ValueError: If an unsupported search mode is specified.
        """
        # Extract query content if it's an Input object
        if hasattr(query, 'content'):
            query = query.content
        if mode == "embedding":
            if not self.has_embedding_search():
                logger.warning("Embedding search not available, falling back to keyword search")
                return self._keyword_search(query, k)
            return self._embedding_search(query, k)
        elif mode == "keyword":
            return self._keyword_search(query, k)
        elif mode == "hybrid":
            return self._hybrid_search(query, k)
        else:
            raise ValueError(f"Unsupported search mode: {mode}")
    
    def _keyword_search(self, query: str, k: int = 5) -> List[Tuple[Document, float]]:
        """
        Search for documents using keyword matching.
        
        Args:
            query: The query to search for
            k: The number of results to return
            
        Returns:
            List[Tuple[Document, float]]: The search results, as (document, score) pairs
        """
        if not TFIDF_AVAILABLE:
            logger.warning("TF-IDF not available, using simple keyword matching")
            return self._simple_keyword_search(query, k)
        
        # Initialize TF-IDF vectorizer if not already done
        if self._tfidf_vectorizer is None or self._tfidf_matrix is None:
            self._tfidf_vectorizer = TfidfVectorizer(stop_words='english')
            documents = [doc.content for doc in self._documents]
            self._tfidf_matrix = self._tfidf_vectorizer.fit_transform(documents)
        
        # Transform query to TF-IDF vector
        query_vector = self._tfidf_vectorizer.transform([query])
        
        # Calculate cosine similarity between query and documents
        from sklearn.metrics.pairwise import cosine_similarity as sklearn_cosine_similarity
        similarities = sklearn_cosine_similarity(query_vector, self._tfidf_matrix).flatten()
        
        # Create (document, score) pairs
        results = [(self._documents[i], float(similarities[i])) for i in range(len(self._documents))]
        
        # Sort by similarity (highest first)
        results.sort(key=lambda x: x[1], reverse=True)
        
        # Return top k results
        return results[:k]
    
    def _simple_keyword_search(self, query: str, k: int = 5) -> List[Tuple[Document, float]]:
        """
        Simple keyword search without using TF-IDF.
        
        Args:
            query: The query to search for
            k: The number of results to return
            
        Returns:
            List[Tuple[Document, float]]: The search results, as (document, score) pairs
        """
        # Normalize query
        query = query.lower()
        query_terms = query.split()
        
        # Calculate scores for each document
        results = []
        for doc in self._documents:
            content = doc.content.lower()
            
            # Count term occurrences
            score = sum(content.count(term) for term in query_terms)
            
            # Normalize by document length
            score = score / (len(content) + 1)  # Add 1 to avoid division by zero
            
            results.append((doc, score))
        
        # Sort by score (highest first)
        results.sort(key=lambda x: x[1], reverse=True)
        
        # Return top k results
        return results[:k]
    
    def _hybrid_search(self, query: str, k: int = 5, 
                      embedding_weight: float = 0.7) -> List[Tuple[Document, float]]:
        """
        Search for documents using both embedding and keyword search.
        
        Args:
            query: The query to search for
            k: The number of results to return
            embedding_weight: Weight to give to embedding search (0-1)
            
        Returns:
            List[Tuple[Document, float]]: The search results, as (document, score) pairs
        """
        # Check if embedding search is available
        if not self.has_embedding_search():
            logger.warning("Embedding search not available, falling back to keyword search")
            return self._keyword_search(query, k)
        
        # Get results from both methods
        embedding_results = self._embedding_search(query, len(self._documents))
        keyword_results = self._keyword_search(query, len(self._documents))
        
        # Create a mapping of document ID to scores
        embedding_scores = {doc.id: score for doc, score in embedding_results}
        keyword_scores = {doc.id: score for doc, score in keyword_results}
        
        # Combine scores
        combined_scores = {}
        all_doc_ids = set(embedding_scores.keys()) | set(keyword_scores.keys())
        
        for doc_id in all_doc_ids:
            emb_score = embedding_scores.get(doc_id, 0.0)
            kw_score = keyword_scores.get(doc_id, 0.0)
            
            # Weighted combination
            combined_scores[doc_id] = (embedding_weight * emb_score) + ((1 - embedding_weight) * kw_score)
        
        # Create (document, score) pairs
        results = []
        for doc_id, score in combined_scores.items():
            doc = self.get_document(doc_id)
            if doc:
                results.append((doc, score))
        
        # Sort by combined score (highest first)
        results.sort(key=lambda x: x[1], reverse=True)
        
        # Return top k results
        return results[:k]
    
    def has_embedding_search(self) -> bool:
        """
        Check if embedding search is available.
        
        Returns:
            bool: True if embedding search is available
        """
        return self._embeddings_cache is not None or self._client is not None
    
    def _embedding_search(self, query: str, k: int = 5) -> List[Tuple[Document, float]]:
        """
        Search for documents similar to the query using embeddings.
        
        Args:
            query: The query to search for.
            k: The number of results to return.
            
        Returns:
            List[Tuple[Document, float]]: The search results, as (document, score) pairs.
        """
        # Get query embedding
        query_embedding = self._client.get_embeddings(query)[0]
        
        # Generate embeddings for documents if they don't exist
        if self._embeddings_cache is None:
            _ = self.embeddings  # This will generate and cache embeddings
        
        similarities = []
        for doc in self._documents:
            if doc.embeddings is None:
                # Skip documents without embeddings
                continue
            
            # Calculate cosine similarity
            similarity = cosine_similarity(np.array(query_embedding), np.array(doc.embeddings))
            similarities.append((doc, similarity))
        
        # Sort by similarity (highest first)
        similarities.sort(key=lambda x: x[1], reverse=True)
        
        # Return top k results
        return similarities[:k]
    
    def has_keyword_search(self) -> bool:
        """
        Check if keyword search is available.
        
        Returns:
            bool: True if keyword search is available
        """
        return True  # Basic keyword search is always available
    
    def __len__(self) -> int:
        """
        Get the number of documents in the knowledge base.
        
        Returns:
            int: The number of documents.
        """
        return len(self._documents)
    
    def __getitem__(self, doc_id: str) -> Document:
        """
        Get a document by ID.
        
        Args:
            doc_id: The ID of the document to get.
            
        Returns:
            Document: The document.
            
        Raises:
            KeyError: If the document is not found.
        """
        doc = self.get_document(doc_id)
        if doc is None:
            raise KeyError(f"Document with ID {doc_id} not found")
        return doc
